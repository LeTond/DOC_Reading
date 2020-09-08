"""
Здесь будет алгоритм для автоматической выгрузки шаблонов и заключений в базу данных
"""
from selenium.webdriver import Chrome
from selenium.webdriver.chrome.options import Options
from docx import Document
from time import sleep
import os

def export(link, root):
    opts = Options()
    opts.set_headless()
    assert opts.headless
    file = open(root + '/log.txt', 'a')
    for fltr in os.listdir(root + '/МРТ/Ангиография/Голова (артерии)/Извитость, изгибы'):
        try:
            if ".docx" in fltr:
                full_text1 = []
                full_text2 = []
                path2 = root + '/МРТ/Ангиография/Голова (артерии)/Извитость, изгибы/' + fltr
                document = Document(path2)
                full_text1.append(document.paragraphs[0].text.lstrip())
                for i in range(1, len(document.paragraphs)):
                    full_text2.append("" + document.paragraphs[i].text.lstrip() + '\n')
                browser = Chrome('/home/lg/Downloads/chromedriver')
                browser.get(link)
                input_form = browser.find_element_by_name('importtool_text')
                input_form.send_keys(full_text1 + list('\n\n') + full_text2)
                input_form.submit()
                results = browser.find_element_by_xpath("//*[contains(@style,'font-family')]")
                file.write(results.text + "\n\n")
                browser.close()
        except ValueError:
            print(f"Необработанный документ: {fltr}")
        else:
            pass
    file.close()
    quit()




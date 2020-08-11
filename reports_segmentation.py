"""
Сегментируем МРТ, КТ, Рентген исследования.
Создаем два каталога и сохраняем в них рассортированные и обработанные заключения.
Удаляем шапку заключения, ФИО пациента, Дату рождения, область исследования, ФИО Врача
"""

from docx import Document
import os
import docx
import shutil


def create_folder_area(path_: str, root: str):
    new_path = root + '/' + path_.replace('//', '/')
    try:
        os.makedirs(new_path)
    except FileExistsError:
        print(f"Такая папка существует: {new_path}")
    else:
        pass


def create_folder_pathology(path_: str, root: str, key_name: str):
    new_path = root + '/' + path_.replace('//', '/') + '/' + key_name
    try:
        os.makedirs(new_path)
        os.mkdir(root + '/' + path_.replace('//', '/') + '/Прочее')
    except FileExistsError:
        print(f"")
    else:
        pass


def start_segmentation(fltr: str, root: str, path: str, key_for_area: list, path_: str):
    set_list_list = []
    control_set = [k for k in range(len(key_for_area))]
    new_path = root + '/' + path_.replace('//', '/')
    try:
        # Отфильтровывает документы с расширением .docx
        if ".docx" in fltr:
            path2 = root + path + '/' + fltr
            conclusion = Document(path2)
            # Фильтруем текст по ключемым словам
            for para in conclusion.paragraphs:
                # Фильтрация по ключу
                for i in range(len(key_for_area)):
                    for j in range(len(key_for_area[i])):
                        if key_for_area[i][j] in para.text or key_for_area[i][j].capitalize() in para.text:
                            set_list_list.append(i)
                            break
                if set(set_list_list) == set(control_set):
                    new_current_path = new_path + '/' + '(new)' + fltr
                    shutil.copy(path2, new_current_path)
    except ValueError:
        print(f"Необработанный документ: {fltr}")
    else:
        pass


def contin_segmentation(fltr: str, root: str, key_words_for_remove: tuple, key_head_words: tuple,
                        key_for_: list, path_: str, key_name: str, corn_path: str):
    set_list_list = []
    control_set = [k for k in range(len(key_for_))]
    curr_path = root + '/' + corn_path.replace('//', '/')
    new_path = root + '/' + path_.replace('//', '/') + '/' + key_name
    try:
        # Отфильтровывает документы с расширением .docx
        if ".docx" in fltr:
            path2 = curr_path + '/' + fltr
            conclusion = Document(path2)
            # Фильтруем текст по ключемым словам
            for para in conclusion.paragraphs:
                # Фильтрация по ключу
                for i in range(len(key_for_)):
                    for j in range(len(key_for_[i])):
                        if key_for_[i][j] in para.text or key_for_[i][j].capitalize() in para.text:
                            set_list_list.append(i)
                            break
            new_doc_step1 = docx.Document()
            if set(set_list_list) == set(control_set):
                for parag_ in range(len(conclusion.paragraphs)):
                    # Отфильтровываем всё до параграфа со словом пациент ...
                    for key_word in key_head_words:
                        if key_word in conclusion.paragraphs[parag_].text \
                                or key_word.capitalize() in conclusion.paragraphs[parag_].text:
                            for index_key_word in range(parag_):
                                conclusion.paragraphs[index_key_word].text = None
                    # Отфильтровываем, параграфы со словами врач и тд.....
                    for remove_key in key_words_for_remove:
                        if remove_key in conclusion.paragraphs[parag_].text \
                                or remove_key.capitalize() in conclusion.paragraphs[parag_].text:
                            conclusion.paragraphs[parag_].text = None
                # Вставляем путь для области исследования
                text = f"{path_} \n {key_name} \n -report-text-below-"
                new_doc_step1.add_paragraph(text)
                # Создаем новый текстовый документ на базе предыдущего без включения пустых параграфов
                for pg in range(len(conclusion.paragraphs)):
                    if conclusion.paragraphs[pg].text == "":
                        pass
                    else:
                        new_doc_step1.add_paragraph(conclusion.paragraphs[pg].text)
                    new_doc_step1.save(new_path + '/' + fltr)
    except ValueError:
        print(f"Необработанный документ: {fltr}")
    else:
        pass


def remove_segmented(fltr: str, root: str, path_: str, key_name: str, corn_path: str):
    curr_path = root + '/' + corn_path.replace('//', '/')
    new_path = root + '/' + path_.replace('//', '/') + '/' + key_name
    if fltr in os.listdir(new_path):
        os.remove(curr_path + '/' + fltr)


def segment_else(fltr: str, root: str, key_words_for_remove: tuple, key_head_words: tuple, path_: str):
    curr_path = root + '/' + path_.replace('//', '/')
    path_else = curr_path + '/Прочее'
    path2 = curr_path + '/' + fltr
    # Завершающий этап сортировки.........костыль
    try:
        # Отфильтровывает документы с расширением .docx
        if ".docx" in fltr:
            conclusion = Document(path2)
            new_doc_step2 = docx.Document()
            # Фильтруем текст по ключемым словам

            for parag_ in range(len(conclusion.paragraphs)):
                # Отфильтровываем всё до параграфа со словом пациент ...
                for key_word in key_head_words:
                    if key_word in conclusion.paragraphs[parag_].text \
                            or key_word.capitalize() in conclusion.paragraphs[parag_].text:
                        for index_key_word in range(parag_):
                            conclusion.paragraphs[index_key_word].text = None
                # Отфильтровываем, параграфы со словами врач и тд.....
                for remove_key in key_words_for_remove:
                    if remove_key in conclusion.paragraphs[parag_].text \
                            or remove_key.capitalize() in conclusion.paragraphs[parag_].text:
                        conclusion.paragraphs[parag_].text = None

            # Вставляем путь для области исследования
            text = f"{path_} \n Прочее \n -report-text-below-"
            new_doc_step2.add_paragraph(text)
            # Создаем новый текстовый документ на базе предыдущего без включения пустых параграфов
            for pg in range(len(conclusion.paragraphs)):
                if conclusion.paragraphs[pg].text == "":
                    pass
                else:
                    new_doc_step2.add_paragraph(conclusion.paragraphs[pg].text)
            # Сохраняем docx со старым названием + (new) по новому пути в папку
            new_doc_step2.save(path_else + '/' + fltr)
            os.remove(curr_path + '/' + fltr)
    except ValueError:
        print(f"Необработанный документ: {fltr}")
    else:
        pass
